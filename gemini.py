import os
import json
import ast
import base64
import tempfile
import re
import logging
import subprocess
import requests
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple
import torch
from transformers import TrOCRProcessor, VisionEncoderDecoderModel
import fitz  
from PIL import Image, ImageEnhance
import cv2
import numpy as np
from fastapi import FastAPI, UploadFile, File, HTTPException, Query
from fastapi.responses import JSONResponse
import uuid
import sqlite3
import smtplib
import threading
import time
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import uuid
from fastapi import Form
from fastapi.responses import HTMLResponse
import traceback

# Import from our new modules
try:
    from email_monitor import start_email_monitoring, stop_email_monitoring
    from approval_system import init_database, save_for_approval, send_approval_email, get_approval_status, update_approval_status
except ImportError:
    # Define dummy functions if modules not found
    def start_email_monitoring(): return False
    def stop_email_monitoring(): pass
    def init_database(): pass
    def save_for_approval(*args, **kwargs): return "test123"
    def send_approval_email(*args, **kwargs): return False
    def get_approval_status(*args, **kwargs): return None
    def update_approval_status(*args, **kwargs): pass
 
try:
    from config import SERVER_HOST, SERVER_PORT, SERVER_URL
except ImportError:
    # Fallback values if config.py doesn't exist
    SERVER_HOST = "0.0.0.0"
    SERVER_PORT = 3002
    SERVER_URL = "http://localhost:3002"

# Force a usable temp directory
CUSTOM_TMP = os.path.expanduser("~/tmp_port_enhanced")
os.makedirs(CUSTOM_TMP, exist_ok=True)
tempfile.tempdir = CUSTOM_TMP
os.environ["TMPDIR"] = os.environ["TEMP"] = os.environ["TMP"] = CUSTOM_TMP
 
# Word docs
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False
 
# textract fallback for .doc
try:
    import textract
    TEXTRACT_AVAILABLE = True
except Exception:
    TEXTRACT_AVAILABLE = False
 
# Tesseract (optional) - FOR MULTI-LANGUAGE
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
    # Check if Arabic language is available
    try:
        pytesseract.get_languages(config='')
        ARABIC_AVAILABLE = 'ara' in pytesseract.get_languages(config='')
        log.info(f"Tesseract languages available: {pytesseract.get_languages(config='')}")
        log.info(f"Arabic support: {ARABIC_AVAILABLE}")
    except:
        ARABIC_AVAILABLE = False
        log.warning("Could not check Tesseract languages")
except Exception:
    TESSERACT_AVAILABLE = False
    pytesseract = None
    ARABIC_AVAILABLE = False
 
# pdf2image (optional)
try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except Exception:
    PDF2IMAGE_AVAILABLE = False
 
# Gemini SDK (Google genai) - For extraction
try:
    from google import genai
    GEMINI_AVAILABLE = True
except Exception:
    GEMINI_AVAILABLE = False
    genai = None
 
# Requests for SAP posting
import requests
 
# Logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s",
)
log = logging.getLogger("combined-extractor")
 
app = FastAPI(title="Unified Document Extractor", version="1.0.0")
 
# ---------------- Target schemas (invoice & certificate) ----------------
INVOICE_SCHEMA: Dict[str, Any] = {
    "invoice_number": "",
    "invoice_date": "",
    "supplier": {"name": "", "vat_number": "", "address": "", "phone": "", "email": ""},
    "buyer": {"name": "", "vat_number": "", "address": "", "phone": "", "email": ""},
    "customer_reference": "",
    "sales_order": "",
    "due_date": "",
    "currency": "",
    "line_items": [
        {"item_no": "", "material_code": "", "description": "", "hs_code": "", "quantity": "",
         "uom": "", "unit_price": "", "net_price": ""}
    ],
    "net_total": "",
    "total_payable": "",
    "payment_terms": "",
    "bank_details": {"beneficiary": "", "bank": "", "iban": "", "account_no": "", "swift": ""},
    "shipment": {"incoterms": "", "location": "", "country_origin": ""},
    "production_date": "",
    "expiry_date": "",
    "notes": "",
    "extras": {}
}
 
CERTIFICATE_SCHEMA: Dict[str, Any] = {
    "certificate_number": "",
    "issue_date": "",
    "certificate_type": "",
    "exporter": {"name": "", "address": ""},
    "consignee": {"name": "", "address": ""},
    "product": {"description": "", "hs_code": "", "lot_number": "", "production_date": "", "expiry_date": ""},
    "quantity": {"net_weight": "", "gross_weight": "", "packaging": ""},
    "shipment": {"date": "", "port": "", "vessel": ""},
    "fob_value": "",
    "issuing_authority": "",
    "extras": {}
}
 
# ---------- SAP OData config ----------
SAP_ODATA_BASE_URL = os.getenv("SAP_ODATA_BASE_URL", "https://103.194.242.134:8005/sap/opu/odata/sap/ZTALAL_INVC_CDS")
SAP_ENTITY = os.getenv("SAP_ENTITY", "ZTALAL_Invc")
SAP_USER = os.getenv("SAP_HTTP_USER", "Vasanthasri")
SAP_PASS = os.getenv("SAP_HTTP_PASS", "Sri@2004")
SAP_VERIFY_SSL = False
SAP_CLIENT = os.getenv("SAP_CLIENT", "800")
 
GEMINI_API_KEY = os.getenv(
    "GEMINI_API_KEY",
    "AIzaSyBFxdoZCxOr2k54wjiYXzpYSS_PVp8SJkY"
).strip()
 
# ---------- Gemini config ----------
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.0-flash-exp")  # Using experimental model for better performance
 
_gemini_client: Optional[genai.Client] = None
 
def get_gemini_client() -> genai.Client:
    global _gemini_client
    if _gemini_client is not None:
        return _gemini_client
    if not GEMINI_AVAILABLE:
        raise RuntimeError("Gemini SDK not installed/available")
    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY not set")
    _gemini_client = genai.Client(api_key=GEMINI_API_KEY)
    return _gemini_client
 
# ------------------ Utility functions ------------------
def pick_device() -> str:
    if torch.cuda.is_available():
        return "cuda"
    if torch.backends.mps.is_available():
        return "mps"
    return "cpu"
 
def safe_remove(path: str):
    try:
        os.remove(path)
    except Exception:
        pass
 
# ------------------ Enhanced Text extraction helpers ------------------
def extract_text_from_docx(path: str) -> str:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed (pip install python-docx)")
    doc = DocxDocument(path)
    paragraphs = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            paragraphs.append(row_text)
    return "\n".join(paragraphs).strip()
 
def extract_text_from_doc(path: str) -> str:
    if TEXTRACT_AVAILABLE:
        try:
            text = textract.process(path).decode("utf-8", errors="ignore")
            return text
        except Exception as e:
            log.warning(f"textract failed for .doc file: {e}")
    try:
        out = subprocess.check_output(["antiword", path], stderr=subprocess.STDOUT, timeout=15)
        return out.decode("utf-8", errors="ignore")
    except Exception as e:
        log.warning(f"antiword fallback failed: {e}")
    return ""
 
def pdf_to_text_native(pdf_path: str) -> str:
    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        log.warning(f"PyMuPDF open failed: {e}")
        return ""
    if doc.page_count == 0:
        return ""
    pages_text = []
    for i in range(doc.page_count):
        try:
            page = doc.load_page(i)
            txt = page.get_text("text") or ""
            pages_text.append(txt)
        except Exception as e:
            log.debug(f"page {i+1} get_text failed: {e}")
    return "\n\n".join(pages_text).strip()
 
# Enhanced preprocessing for handwritten text
def preprocess_image_for_handwritten_ocr(pil_image: Image.Image) -> Image.Image:
    """
    Enhanced preprocessing specifically for handwritten invoices
    """
    try:
        img_np = np.array(pil_image.convert("RGB"))
        
        # Convert to grayscale
        gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
        
        # Apply CLAHE (Contrast Limited Adaptive Histogram Equalization) for better contrast
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
        gray = clahe.apply(gray)
        
        # Denoise
        denoised = cv2.medianBlur(gray, 3)
        
        # Binarization - try multiple methods
        # Method 1: Adaptive Gaussian
        binary1 = cv2.adaptiveThreshold(denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                       cv2.THRESH_BINARY, 11, 2)
        
        # Method 2: Otsu's threshold
        _, binary2 = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Method 3: Mean adaptive
        binary3 = cv2.adaptiveThreshold(denoised, 255, cv2.ADAPTIVE_THRESH_MEAN_C, 
                                       cv2.THRESH_BINARY, 11, 2)
        
        # Combine results - use the best looking one
        # Count black pixels (text)
        black_pixels1 = np.sum(binary1 == 0)
        black_pixels2 = np.sum(binary2 == 0)
        black_pixels3 = np.sum(binary3 == 0)
        
        # Choose the one with most black pixels (but not too many)
        binary = binary1
        if 5000 < black_pixels2 < 50000 and black_pixels2 > black_pixels1:
            binary = binary2
        if 5000 < black_pixels3 < 50000 and black_pixels3 > black_pixels2:
            binary = binary3
        
        # Apply morphological operations to clean up
        kernel = np.ones((2,2), np.uint8)
        binary = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)
        binary = cv2.morphologyEx(binary, cv2.MORPH_OPEN, kernel)
        
        # Sharpen
        kernel = np.array([[-1,-1,-1], [-1,9,-1], [-1,-1,-1]])
        sharpened = cv2.filter2D(binary, -1, kernel)
        
        # Convert back to PIL
        pil = Image.fromarray(sharpened)
        
        # Enhance contrast and brightness
        enhancer = ImageEnhance.Contrast(pil)
        pil = enhancer.enhance(2.0)
        
        enhancer2 = ImageEnhance.Brightness(pil)
        pil = enhancer2.enhance(1.2)
        
        return pil.convert("RGB")
        
    except Exception as e:
        log.error(f"Error in handwritten preprocessing: {e}")
        return pil_image

# Add near other utility functions
def split_pdf_by_pages(pdf_path: str, output_dir: str) -> List[str]:
    """
    Split a PDF into individual pages
    Returns list of file paths for each page
    """
    try:
        doc = fitz.open(pdf_path)
        page_files = []
        
        for page_num in range(doc.page_count):
            # Create a new PDF with only this page
            output_pdf = fitz.open()
            output_pdf.insert_pdf(doc, from_page=page_num, to_page=page_num)
            
            # Save the single page PDF
            page_filename = f"page_{page_num+1}_{os.path.basename(pdf_path)}"
            page_path = os.path.join(output_dir, page_filename)
            output_pdf.save(page_path)
            output_pdf.close()
            
            page_files.append(page_path)
            
        doc.close()
        log.info(f"Split PDF into {len(page_files)} pages")
        return page_files
        
    except Exception as e:
        log.error(f"Error splitting PDF: {e}")
        return []

def process_single_page_invoice(pdf_path: str, page_num: int = 1, 
                                email_subject: str = "", email_from: str = "") -> Optional[str]:
    """
    Process a single page PDF as an invoice
    Returns approval_id if successful
    """
    try:
        log.info(f"Processing page {page_num}: {os.path.basename(pdf_path)}")
        
        # Extract data using Gemini PDF
        gemini_invoices = extract_invoice_with_gemini_pdf(pdf_path)
        
        if not gemini_invoices:
            log.warning(f"No invoice data extracted from page {page_num}")
            return None
        
        # Use the first (and should be only) invoice from this page
        extracted_data = normalize_gemini_pdf_invoice(gemini_invoices[0])
        
        # Build SAP payload
        sap_payload = build_sap_invoice_payload(extracted_data)
        
        # Save for approval with page info
        approval_id = save_for_approval(
            file_path=pdf_path,
            extracted_data=extracted_data,
            sap_payload=sap_payload,
            email_subject=f"{email_subject} (Page {page_num})" if email_subject else f"Page {page_num}"
        )
        
        if approval_id:
            if send_approval_email(approval_id, extracted_data):
                log.info(f"✅ Approval email sent for page {page_num} invoice: {approval_id}")
            else:
                log.error(f"❌ Failed to send approval email for page {page_num}")
        else:
            log.error(f"Failed to save for approval for page {page_num}")
            
        return approval_id
        
    except Exception as e:
        log.error(f"Error processing page {page_num}: {e}")
        log.error(traceback.format_exc())
        return None

def extract_invoice_with_gemini_pdf(pdf_path: str) -> List[Dict[str, Any]]:
    """
    Uses Gemini 2.0 Flash to extract invoice data DIRECTLY from PDF
    Enhanced to handle single-page invoices better
    """
    client = get_gemini_client()

    with open(pdf_path, "rb") as f:
        pdf_bytes = f.read()

    pdf_b64 = base64.b64encode(pdf_bytes).decode("utf-8")

    prompt = """
Analyze this PDF containing restaurant invoices.
This PDF contains ONE invoice per page. Extract invoice data into a structured JSON array.

CRITICAL EXTRACTION RULES:
1. Invoice Number: Look for "Invoice No:", "Inv No:", "#", "Number:", "Invoice #"
2. Supplier/Vendor Name: Look at the TOP of the document for business names. 
   Common patterns: "Company Name:", "Supplier:", "Vendor:", "Issued by:", "From:"
   Look for business suffixes: TRADING, LLC, FZE, EST, GROUP
3. Date: Look for "Date:", "Invoice Date:", "Dated:", "Issued:"
4. Customer: Look for "Bill To:", "Customer:", "Ship To:", "Name:"
5. Line Items: Extract ALL items from any table/listing
6. Totals:
   - sub_total
   - tax_amount
   - grand_total

IMPORTANT FOR SUPPLIER NAME:
- Extract the complete business name from the top of the invoice
- Include the full legal name (e.g., "AL RAHMAN TRADING LLC", "BIO GREEN FZE")
- Don't shorten or abbreviate unnecessarily

IMPORTANT FOR NUMBERS:
IMPORTANT FOR UAE HANDWRITTEN TOTALS (DHS / FILS TWO-COLUMN):
- Many invoices show amounts in 2 columns: "Dhs" and "Fils".
- Example: Dhs=34 and Fils=65 means 34.65
- If totals/VAT appear split like "34 65" or in two adjacent columns, you MUST output as 34.65
- Same for VAT: Dhs=1 and Fils=65 => 1.65
- Convert European format (1.234,56) to standard format (1234.56)
- Remove thousand separators (dots in European format)
- Use dot (.) as decimal separator

RETURN ONLY VALID JSON ARRAY with this structure:
[
  {
    "invoice_number": "value",
    "supplier_name": "value",  // FULL business name
    "vendor_name": "value",
    "date": "value",
    "customer": "value",
    "line_items": [
      {
        "description": "value",
        "quantity": "value",
        "unit_price": "value",
        "total": "value"
      }
    ],
    "sub_total": "value",
    "tax_amount": "value",
    "grand_total": "value",
    "currency": "value"
  }
]

If a field is not found, use empty string "".
"""

    response = client.models.generate_content(
        model="gemini-2.0-flash",
        contents=[
            {
                "role": "user",
                "parts": [
                    {"text": prompt},
                    {
                        "inline_data": {
                            "mime_type": "application/pdf",
                            "data": pdf_b64
                        }
                    }
                ]
            }
        ],
        config={
            "temperature": 0.1,
            "max_output_tokens": 4000,
            "response_mime_type": "application/json"
        }
    )

    raw = getattr(response, "text", "") or ""
    if not raw:
        raise RuntimeError("Empty Gemini PDF response")

    try:
        return json.loads(raw)
    except Exception as e:
        log.error(f"Failed to parse Gemini PDF JSON: {e}")
        log.error(raw[:500])
        return []

def normalize_gemini_pdf_invoice(gemini_inv: Dict[str, Any]) -> Dict[str, Any]:
    """
    Converts Gemini PDF invoice output into internal invoice schema
    """
    # Improved supplier name extraction from Gemini response
    vendor_name = str(gemini_inv.get("vendor_name", ""))
    supplier_name = str(gemini_inv.get("supplier_name", ""))
    
    # Try multiple possible keys for supplier name
    supplier_keys = ["supplier", "vendor", "seller", "company_name", "business_name", "issuer"]
    
    # Look for supplier name in various possible locations
    for key in supplier_keys:
        if key in gemini_inv:
            if isinstance(gemini_inv[key], dict) and "name" in gemini_inv[key]:
                supplier_name = gemini_inv[key]["name"]
                break
            elif isinstance(gemini_inv[key], str):
                supplier_name = gemini_inv[key]
                break
    
    # If still empty, try to find from text patterns
    if not supplier_name and not vendor_name:
        # Check if Gemini returned raw text with supplier info
        raw_text = gemini_inv.get("extracted_text", "")
        if raw_text:
            # Look for company name patterns in the beginning of document
            lines = raw_text.split('\n')
            for line in lines[:10]:  # Check first 10 lines (where supplier usually is)
                line = line.strip()
                # Common business suffixes in UAE
                business_suffixes = ['TRADING', 'LLC', 'FZE', 'L.L.C', 'EST', 'GROUP']
                if any(suffix in line.upper() for suffix in business_suffixes):
                    supplier_name = line
                    break
                # Also check for lines in all caps (common for company headers)
                elif line.isupper() and len(line) > 3 and len(line) < 100:
                    supplier_name = line
                    break
    
    # Use whichever is available
    final_supplier_name = supplier_name or vendor_name
    
    # Clean up the supplier name
    if final_supplier_name:
        # Remove common unwanted prefixes/suffixes
        final_supplier_name = re.sub(r'^(Supplier|Vendor|Company|Name)[:\s]*', '', final_supplier_name, flags=re.IGNORECASE)
        final_supplier_name = final_supplier_name.strip()
    
    # Process line items
    line_items = gemini_inv.get("line_items", [])
    if isinstance(line_items, list):
        # Ensure each line item has required structure
        processed_items = []
        for item in line_items:
            if isinstance(item, dict):
                processed_items.append({
                    "item_no": str(item.get("item_no", "")),
                    "material_code": str(item.get("material_code", "")),
                    "description": str(item.get("description", item.get("name", ""))),
                    "hs_code": str(item.get("hs_code", "")),
                    "quantity": str(item.get("quantity", "")),
                    "uom": str(item.get("uom", item.get("unit", ""))),
                    "unit_price": str(item.get("unit_price", item.get("price", ""))),
                    "net_price": str(item.get("net_price", item.get("total", "")))
                })
        line_items = processed_items
    
    return {
        "invoice_number": str(gemini_inv.get("invoice_number", gemini_inv.get("invoice_no", ""))),
        "invoice_date": str(gemini_inv.get("date", gemini_inv.get("invoice_date", ""))),
        "supplier": {
            "name": final_supplier_name,
            "vat_number": str(gemini_inv.get("vat_number", "")),
            "address": str(gemini_inv.get("address", gemini_inv.get("supplier_address", ""))),
            "phone": str(gemini_inv.get("phone", gemini_inv.get("supplier_phone", ""))),
            "email": str(gemini_inv.get("email", gemini_inv.get("supplier_email", "")))
        },
        "buyer": gemini_inv.get("buyer", {"name": "", "vat_number": "", "address": "", "phone": "", "email": ""}),
        "customer_reference": str(gemini_inv.get("customer_reference", "")),
        "sales_order": str(gemini_inv.get("sales_order", "")),
        "due_date": str(gemini_inv.get("due_date", "")),
        "currency": str(gemini_inv.get("currency", "AED")),
        "line_items": line_items,
        "net_total": str(gemini_inv.get("net_total", gemini_inv.get("sub_total", ""))),
        "total_payable": str(gemini_inv.get("grand_total", gemini_inv.get("total_payable", ""))),
        "payment_terms": str(gemini_inv.get("payment_terms", "")),
        "bank_details": gemini_inv.get("bank_details", {"beneficiary": "", "bank": "", "iban": "", "account_no": "", "swift": ""}),
        "shipment": gemini_inv.get("shipment", {"incoterms": "", "location": "", "country_origin": ""}),
        "production_date": str(gemini_inv.get("production_date", "")),
        "expiry_date": str(gemini_inv.get("expiry_date", "")),
        "notes": str(gemini_inv.get("notes", "")),
        "extras": {
            "source": "gemini_pdf",
            "raw_gemini_response": gemini_inv
        }
    }

 
def ocr_with_tesseract_image_enhanced(image: Image.Image) -> str:
    """
    Enhanced Tesseract OCR with multi-language support for Arabic/English
    """
    if not TESSERACT_AVAILABLE or pytesseract is None:
        return ""
    
    try:
        img_np = np.array(image.convert("RGB"))
        gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
        
        # Try different preprocessing for better results
        results = []
        
        # Method 1: Otsu's threshold
        _, thresh1 = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        pil1 = Image.fromarray(thresh1)
        
        # Method 2: Adaptive threshold
        thresh2 = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                       cv2.THRESH_BINARY, 11, 2)
        pil2 = Image.fromarray(thresh2)
        
        # Try with English first
        config_eng = r'--oem 3 --psm 6 -l eng'
        
        # Try with Arabic+English if available
        if ARABIC_AVAILABLE:
            config_ara_eng = r'--oem 3 --psm 6 -l ara+eng'
        else:
            config_ara_eng = config_eng
        
        # Try multiple configurations
        configs_to_try = [
            config_eng,
            config_ara_eng,
            r'--oem 3 --psm 4 -l eng',  # Different page segmentation
            r'--oem 3 --psm 11 -l eng',  # Sparse text
        ]
        
        best_text = ""
        best_config = ""
        
        for config in configs_to_try:
            try:
                # Try with first image
                text1 = pytesseract.image_to_string(pil1, config=config).strip()
                
                # Try with second image
                text2 = pytesseract.image_to_string(pil2, config=config).strip()
                
                # Choose the better result
                text = text1 if len(text1) > len(text2) else text2
                
                if len(text) > len(best_text):
                    best_text = text
                    best_config = config
                    
            except Exception as e:
                log.debug(f"Tesseract config {config} failed: {e}")
                continue
        
        log.info(f"Best Tesseract config: {best_config}, chars extracted: {len(best_text)}")
        return best_text
        
    except Exception as e:
        log.warning(f"Enhanced Tesseract OCR failed: {e}")
        return ""
 
_trocr_printed_proc = None
_trocr_printed_model = None
_trocr_printed_device = None

_trocr_handwritten_proc = None
_trocr_handwritten_model = None
_trocr_handwritten_device = None
 
def ocr_with_trocr_printed_image(image: Image.Image) -> str:
    """TrOCR for printed text"""
    global _trocr_printed_proc, _trocr_printed_model, _trocr_printed_device
    try:
        if _trocr_printed_proc is None or _trocr_printed_model is None:
            _trocr_printed_proc = TrOCRProcessor.from_pretrained("microsoft/trocr-large-printed")
            _trocr_printed_model = VisionEncoderDecoderModel.from_pretrained("microsoft/trocr-large-printed")
            _trocr_printed_device = pick_device()
            _trocr_printed_model.to(_trocr_printed_device).eval()
        
        img_pre = preprocess_image_for_handwritten_ocr(image)
        inputs = _trocr_printed_proc(images=img_pre, return_tensors="pt").to(_trocr_printed_device)
        with torch.no_grad():
            out_ids = _trocr_printed_model.generate(**inputs, max_length=2048)
        decoded = _trocr_printed_proc.batch_decode(out_ids, skip_special_tokens=True)[0]
        return decoded.strip()
    except Exception as e:
        log.warning(f"TrOCR printed failed: {e}")
        return ""
 
def ocr_with_trocr_handwritten_image(image: Image.Image) -> str:
    """TrOCR for handwritten text"""
    global _trocr_handwritten_proc, _trocr_handwritten_model, _trocr_handwritten_device
    try:
        if _trocr_handwritten_proc is None or _trocr_handwritten_model is None:
            _trocr_handwritten_proc = TrOCRProcessor.from_pretrained("microsoft/trocr-base-handwritten")
            _trocr_handwritten_model = VisionEncoderDecoderModel.from_pretrained("microsoft/trocr-base-handwritten")
            _trocr_handwritten_device = pick_device()
            _trocr_handwritten_model.to(_trocr_handwritten_device).eval()
        
        img_pre = preprocess_image_for_handwritten_ocr(image)
        inputs = _trocr_handwritten_proc(images=img_pre, return_tensors="pt").to(_trocr_handwritten_device)
        with torch.no_grad():
            out_ids = _trocr_handwritten_model.generate(**inputs, max_length=2048)
        decoded = _trocr_handwritten_proc.batch_decode(out_ids, skip_special_tokens=True)[0]
        return decoded.strip()
    except Exception as e:
        log.warning(f"TrOCR handwritten failed: {e}")
        return ""
 
def enhanced_ocr_pdf_to_text(pdf_path: str) -> Tuple[str, Dict[str, Any]]:
    """
    Enhanced OCR with multiple engines for better accuracy
    """
    debug_info = {
        "engine_used_by_page": {}, 
        "page_count": 0, 
        "tesseract_available": TESSERACT_AVAILABLE,
        "arabic_support": ARABIC_AVAILABLE
    }
    
    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        return "", {"error": f"fitz open failed: {e}"}
    
    page_count = doc.page_count
    debug_info["page_count"] = page_count
    all_text_pages = []
    
    for i in range(page_count):
        page = doc.load_page(i)
        zoom = 4.0  # Increased zoom for handwritten text
        pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False, dpi=300)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        
        # Store all OCR results
        ocr_results = []
        
        # Method 1: Tesseract with enhanced preprocessing
        if TESSERACT_AVAILABLE:
            tesseract_text = ocr_with_tesseract_image_enhanced(img)
            if tesseract_text and len(tesseract_text.strip()) > 10:
                ocr_results.append(("tesseract", tesseract_text))
                log.info(f"Page {i+1}: Tesseract extracted {len(tesseract_text)} chars")
        
        # Method 2: TrOCR for handwritten
        trocr_handwritten_text = ocr_with_trocr_handwritten_image(img)
        if trocr_handwritten_text and len(trocr_handwritten_text.strip()) > 10:
            ocr_results.append(("trocr_handwritten", trocr_handwritten_text))
            log.info(f"Page {i+1}: TrOCR Handwritten extracted {len(trocr_handwritten_text)} chars")
        
        # Method 3: TrOCR for printed (fallback)
        trocr_printed_text = ocr_with_trocr_printed_image(img)
        if trocr_printed_text and len(trocr_printed_text.strip()) > 10:
            ocr_results.append(("trocr_printed", trocr_printed_text))
            log.info(f"Page {i+1}: TrOCR Printed extracted {len(trocr_printed_text)} chars")
        
        # Choose the best result (longest text with good quality)
        best_text = ""
        best_engine = "none"
        
        for engine, text in ocr_results:
            # Clean the text
            cleaned_text = re.sub(r'\s+', ' ', text).strip()
            
            # Check if it looks like meaningful text (has numbers and letters)
            has_numbers = bool(re.search(r'\d', cleaned_text))
            has_letters = bool(re.search(r'[a-zA-Z]', cleaned_text))
            
            if has_numbers and has_letters and len(cleaned_text) > len(best_text):
                best_text = cleaned_text
                best_engine = engine
        
        # If no good text found, use the longest one
        if not best_text and ocr_results:
            best_engine, best_text = max(ocr_results, key=lambda x: len(x[1]))
        
        debug_info["engine_used_by_page"][f"page_{i+1}"] = best_engine
        all_text_pages.append(f"=== PAGE {i+1} (Engine: {best_engine}) ===\n{best_text}")
    
    combined_text = "\n\n".join(all_text_pages)
    log.info(f"Total OCR text extracted: {len(combined_text)} characters")
    
    return combined_text, debug_info
 
def extract_document_text(pdf_path: str) -> Tuple[str, Dict[str, Any]]:
    """
    Extract text using:
    1. Native PDF text (for digital PDFs)
    2. Fallback → Enhanced OCR (Tesseract/TrOCR with multiple engines)
    """
    debug = {}
    native_text = pdf_to_text_native(pdf_path)
    
    # Check if native text looks like a real document (has invoice-like content)
    if native_text and len(native_text.strip()) > 100:
        # Check for common invoice patterns
        invoice_patterns = [r'invoice', r'bill', r'total', r'amount', r'date', r'no\.', r'number']
        pattern_count = sum(1 for pattern in invoice_patterns if re.search(pattern, native_text.lower()))
        
        if pattern_count >= 2:
            debug["mode"] = "native"
            log.info(f"Using native PDF text with {pattern_count} invoice patterns found")
            return native_text, debug
    
    log.info("Native PDF text insufficient or not found → Using Enhanced OCR")
    ocr_text, ocr_debug = enhanced_ocr_pdf_to_text(pdf_path)
    debug["mode"] = "ocr"
    debug["ocr"] = ocr_debug
    return ocr_text, debug
 
# ------------------ Enhanced Document type detection ------------------
def detect_document_type(text: str) -> str:
    text_lower = (text or "").lower()
    
    # More comprehensive invoice keywords
    invoice_keywords = [
        'invoice', 'bill', 'payment', 'total payable', 'net total', 'vat', 
        'invoice no', 'invoice number', 'inv no', 'inv number',
        'tax invoice', 'bill to', 'amount due', 'subtotal',
        'grand total', 'total amount', 'total aed', 'total dhs'
    ]
    
    invoice_count = sum(1 for k in invoice_keywords if k in text_lower)
    
    cert_keywords = ['certificate', 'certify', 'non-gmo', 'inspection', 'consignee', 'certificate no', 'certificate number']
    cert_count = sum(1 for k in cert_keywords if k in text_lower)
    
    if invoice_count >= 1:  # Lower threshold for handwritten
        return "invoice"
    if cert_count > 2:
        return "certificate"
    
    # Check for invoice number patterns
    invoice_no_patterns = [
        r'invoice\s*(?:no|number|#)?\s*[:]?\s*[\d\-/]+',
        r'inv\s*(?:no|number|#)?\s*[:]?\s*[\d\-/]+',
        r'bill\s*(?:no|number|#)?\s*[:]?\s*[\d\-/]+'
    ]
    
    for pattern in invoice_no_patterns:
        if re.search(pattern, text_lower, re.IGNORECASE):
            return "invoice"
    
    return "unknown"
 
# ------------------ Enhanced Gemini extraction for handwritten ------------------
def parse_gemini_json_like(raw: str) -> Dict[str, Any]:
    raw = (raw or "").strip()
    if not raw:
        raise ValueError("Empty Gemini response")
    
    # First try direct JSON parse
    try:
        return json.loads(raw)
    except Exception:
        pass
    
    # Try to extract JSON from code blocks
    if raw.startswith("```"):
        raw = re.sub(r"^```[a-zA-Z0-9]*\s*", "", raw)
        raw = re.sub(r"```$", "", raw).strip()
    
    # Find JSON object
    start = raw.find("{")
    end = raw.rfind("}")
    if start != -1 and end != -1 and start < end:
        inner = raw[start:end+1]
        try:
            return json.loads(inner)
        except Exception:
            pass
    
    # Try with ast.literal_eval
    try:
        v = ast.literal_eval(raw)
        if isinstance(v, dict):
            return v
    except Exception:
        pass
    
    # If all else fails, return empty schema
    log.warning(f"Could not parse Gemini response: {raw[:200]}...")
    return {}
 
def extract_with_gemini_handwritten(text: str) -> Dict[str, Any]:
    """
    Specialized extraction for handwritten invoices with OCR errors
    """
    if not GEMINI_AVAILABLE:
        return {"error": "Gemini SDK not installed", "raw_text_preview": text[:1000]}
    
    client = get_gemini_client()
    
    # Special prompt for handwritten invoices
    prompt = f"""
You are an expert invoice data extractor specialized in HANDWRITTEN invoices that may contain OCR errors.
The text below comes from OCR of a handwritten invoice in English and/or Arabic.

IMPORTANT INSTRUCTIONS:
1. Extract ONLY these key fields with HIGH ACCURACY:
   - invoice_number: Look for patterns like "Invoice No:", "Inv No:", "No:", "Number:" followed by numbers
   - invoice_date: Look for "Date:", "Invoice Date:", "Dated:" followed by date (could be in DD/MM/YY, DD-MM-YYYY, etc.)
   - supplier.name: Company name at the top of the invoice (look for business names like "TRADING", "LLC", "FZE", etc.)
   - total_payable: Look for "Total AED:", "Total Dhs:", "Grand Total:", "Total Amount:", "Amount Due:" at the END of the invoice
   - currency: Usually "AED" 
   - line_items[0].description: Main product/service description
   - line_items[0].quantity: Quantity (look for numbers after "Qty" or before rate)
   - line_items[0].unit_price: Rate or price per unit

2. Handle OCR errors intelligently:
   - "O" might be "0" (zero), "I" might be "1" (one), "l" might be "1"
   - "BIO GREEN" might appear as "810 GREEN" or similar
   - Numbers might be misread: "28.00" might be "Z8.00" or "28.OO"
   - Dates might be misread: "27/06/25" might be "Z7/06/25"

3. For supplier name: Look at the TOP of the document for business names
4. For total amount: Look at the BOTTOM of the document for final totals
5. If a field is not found, leave it as empty string
6. Return ONLY valid JSON in this exact format:
{{
  "invoice_number": "extracted number or empty",
  "invoice_date": "extracted date or empty",
  "supplier": {{
    "name": "extracted supplier name or empty",
    "vat_number": "",
    "address": "",
    "phone": "",
    "email": ""
  }},
  "total_payable": "extracted total amount or empty",
  "currency": "extracted currency (AED) or empty",
  "line_items": [
    {{
      "item_no": "",
      "material_code": "",
      "description": "main item description or empty",
      "hs_code": "",
      "quantity": "quantity or empty",
      "uom": "",
      "unit_price": "price per unit or empty",
      "net_price": ""
    }}
  ]
}}

7. CLEAN the data:
   - Remove any OCR artifacts like stray characters
   - Convert written numbers to digits
   - Normalize dates to YYYY-MM-DD if possible
   - Remove any text that's clearly noise

OCR TEXT (may contain errors):
{text[:20000]}

EXTRACTED JSON (ONLY JSON, no explanations):
"""
    
    try:
        response = client.models.generate_content(model=GEMINI_MODEL, contents=[prompt])
        raw = getattr(response, "text", "") or ""
        
        if not raw:
            log.warning("Empty response from Gemini")
            return {}
        
        log.info(f"Gemini response length: {len(raw)} chars")
        log.info(f"First 500 chars: {raw[:500]}")
        
        try:
            extracted = parse_gemini_json_like(raw)
            
            # Validate and clean extracted data
            if not isinstance(extracted, dict):
                extracted = {}
            
            # Ensure required structure
            result = {
                "invoice_number": extracted.get("invoice_number", ""),
                "invoice_date": extracted.get("invoice_date", ""),
                "supplier": extracted.get("supplier", {"name": ""}),
                "total_payable": extracted.get("total_payable", ""),
                "currency": extracted.get("currency", ""),
                "line_items": extracted.get("line_items", [{}]),
                "extras": {
                    "source": "handwritten_enhanced",
                    "ocr_text_preview": text[:500]
                }
            }
            
            # Clean up supplier name
            if isinstance(result["supplier"], dict):
                supplier_name = result["supplier"].get("name", "")
                # Remove common OCR errors from supplier name
                supplier_name = re.sub(r'[^\w\s\-\.&,]', '', supplier_name)
                supplier_name = re.sub(r'\s+', ' ', supplier_name).strip()
                result["supplier"]["name"] = supplier_name
            
            # Clean total payable - remove non-numeric except decimal point
            total = result["total_payable"]
            if total:
                # Common OCR errors in numbers
                total = total.replace('O', '0').replace('o', '0')
                total = total.replace('I', '1').replace('l', '1')
                total = total.replace('Z', '2').replace('z', '2')
                total = total.replace('S', '5').replace('s', '5')
                total = total.replace('B', '8').replace('b', '8')
                
                # Extract only numbers and decimal
                total_clean = re.sub(r'[^\d\.]', '', total)
                if total_clean:
                    result["total_payable"] = total_clean
            
            log.info(f"Extracted invoice number: {result['invoice_number']}")
            log.info(f"Extracted total: {result['total_payable']}")
            log.info(f"Extracted currency: {result['currency']}")
            
            return result
            
        except Exception as parse_error:
            log.error(f"Failed to parse Gemini response: {parse_error}")
            log.error(f"Raw response: {raw[:500]}")
            return {}
            
    except Exception as e:
        log.error(f"Gemini call failed: {e}")
        return {}
 
def extract_with_gemini(text: str, doc_type: str) -> Dict[str, Any]:
    """
    Main extraction function with fallback to handwritten specialized extraction
    """
    if not GEMINI_AVAILABLE:
        return {"error": "Gemini SDK not installed", "raw_text_preview": text[:1000]}
    
    # For invoices, use specialized extraction
    if doc_type == "invoice":
        # Check if this looks like handwritten (based on OCR quality indicators)
        ocr_quality_indicators = [
            r'[oO0]{2,}',  # Multiple Os/0s together
            r'[il1|]{2,}',  # Confusion between i, l, 1, |
            r'[zZ2]{2,}',  # Confusion between z, Z, 2
            r'[sS5]{2,}',  # Confusion between s, S, 5
        ]
        
        indicator_count = sum(1 for pattern in ocr_quality_indicators if re.search(pattern, text))
        
        if indicator_count > 2 or len(text) < 500:  # Likely handwritten or poor OCR
            log.info("Detected possible handwritten/poor quality OCR → Using specialized extraction")
            return extract_with_gemini_handwritten(text)
        else:
            # Use standard extraction for printed/clean invoices
            log.info("Using standard Gemini extraction for clean/printed invoice")
            return extract_with_gemini_standard(text, doc_type)
    else:
        # For certificates, use standard extraction
        return extract_with_gemini_standard(text, doc_type)
 
def extract_with_gemini_standard(text: str, doc_type: str) -> Dict[str, Any]:
    """Standard Gemini extraction for clean documents"""
    if doc_type == "invoice":
        schema = INVOICE_SCHEMA
        instructions = "- Extract ALL invoice fields accurately\n- Handle dates, amounts, and numbers carefully\n- Return only JSON matching schema"
    elif doc_type == "certificate":
        schema = CERTIFICATE_SCHEMA
        instructions = "- Extract certificate fields..."
    else:
        return {"raw_text": text, "document_type": doc_type}
 
    schema_json = json.dumps(schema, indent=2)
    prompt = f"""
You are an invoice/certificate extraction engine.
Document type: {doc_type}
INSTRUCTIONS:
1) Extract ALL information from the document text.
2) Return valid JSON matching this schema exactly:
{schema_json}
3) For fields not found leave empty strings. Put any extra fields under "extras".
4) Clean OCR noise and whitespace.
5) Numbers should be normalized (remove thousand separators).
6) Return ONLY the JSON object (no explanation).

DOCUMENT TEXT:
{text[:15000]}
"""
    
    try:
        client = get_gemini_client()
        response = client.models.generate_content(model=GEMINI_MODEL, contents=[prompt])
        raw = getattr(response, "text", "") or ""
        
        if not raw:
            return schema
        
        try:
            return parse_gemini_json_like(raw)
        except Exception:
            m = re.search(r'\{.*\}', raw, re.DOTALL)
            if m:
                try:
                    return json.loads(m.group())
                except Exception:
                    pass
            return schema
    except Exception as e:
        log.error(f"Gemini call failed: {e}")
        return schema
 
# ------------------ SAP helpers (unchanged) ------------------
def convert_date_to_sap_format(date_str: str) -> str:
    if not date_str:
        return ""
    for fmt in ("%d.%m.%Y", "%d-%m-%Y", "%d/%m/%Y", "%Y-%m-%d", "%d/%m/%y", "%d-%m-%y"):
        try:
            dt = datetime.strptime(date_str, fmt)
            return dt.strftime("%Y-%m-%dT00:00:00")
        except Exception:
            continue
    return date_str
 
def safe_float(s: str) -> float:
    try:
        return float(str(s).replace(",", "").strip())
    except Exception:
        return 0.0
 
def ensure_numeric_value(val: str) -> str:
    if val is None:
        return "0"
    val = str(val).strip()
    if val in ("", ".", "-"):
        return "0"
    val = val.replace(",", "")
    try:
        float(val)
        return val
    except Exception:
        return "0"
 
def safe_clean(value: Any) -> str:
    """
    Returns a clean printable string.
    If the value is None, invalid, contains non-printable characters,
    or conversion fails → returns empty string.
    """
    try:
        if value is None:
            return ""
        text = str(value).strip()
        # Remove non-printable characters
        text = ''.join(c for c in text if c.isprintable())
        # If nothing left, return empty
        if not text or text in ["None", "nan", "null", ".", "-", "--"]:
            return ""
        return text
    except Exception:
        return ""

def get_next_ref_num():
    try:
        url = f"{SAP_ODATA_BASE_URL}{SAP_ENTITY}?$orderby=ZRefNum desc&$top=1"
        
        r = requests.get(url, auth=(SAP_USER, SAP_PASS), verify=False, timeout=10)
        r.raise_for_status()
        data = r.json()
        
        results = data.get("d", {}).get("results", []) or data.get("value", [])
        
        if results:
            last_ref = results[0].get("ZRefNum")
            if last_ref and last_ref.startswith("SAP"):
                num_part = last_ref[3:]
                if num_part.isdigit():
                    next_num = int(num_part) + 1
                else:
                    next_num = 1
            else:
                next_num = 1
            
            return f"SAP{next_num:06d}"
    
    except Exception as e:
        log.error(f"Error fetching next ref number: {e}")
        log.error(f"Response text: {r.text if 'r' in locals() else 'No response'}")
        fallback_num = int(uuid.uuid4().int % 1000000)
        return f"SAP{fallback_num:06d}"

def build_sap_invoice_payload(extracted: Dict[str, Any]) -> Dict[str, Any]:
    inv = extracted or {}
    supplier = inv.get("supplier", {}) or {}  
    supplier_name = supplier.get("name", "")
    c_supp = safe_clean(supplier_name)[:50]
    line_items = inv.get("line_items") or []
    first_item = line_items[0] if line_items else {}
    inv_num = inv.get("invoice_number") or "0"
    invoice_date_raw = extracted.get("invoice_date", "")
    invoice_date_sap = convert_date_to_sap_format(invoice_date_raw)
    
    # FIXED: Use AED as default currency
    currency = "AED"
    
    grand_total = inv.get("total_payable") or "0.00"
    tax_total = inv.get("tax_total") or "0.00"
    desc = first_item.get("description") or ""
    clean_desc = re.sub(r"[^A-Za-z0-9\s\-]", "", desc or "")
    clean_desc = clean_desc[:50] 
    uom = first_item.get("uom") or ""
    qty_str = ensure_numeric_value(first_item.get("quantity") or "0")
    qty = safe_float(qty_str)
    unit_price = (first_item.get("unit_price") or "").strip()
    unit_price = ensure_numeric_value(unit_price)
 
    if safe_float(unit_price) == 0 and qty > 0 and safe_float(grand_total) > 0:
        unit_price = f"{safe_float(grand_total) / qty:.2f}"
    unit_price = unit_price or "0.00"
    tax_total = ensure_numeric_value(tax_total)
    
    def format_for_sap(amount_str: str, decimals: int = 2) -> str:
        """Convert any number format to SAP format with fixed decimals."""
        if amount_str is None:
            return "0.00" if decimals == 2 else "0.000"

        s = str(amount_str).strip()
        if not s:
            return "0.00" if decimals == 2 else "0.000"

        # Remove any currency symbols or text, keep digits, comma, dot, minus
        s = re.sub(r"[^\d\.,\-]", "", s)

        # European format: dot as thousand, comma as decimal (2.948,00)
        if "." in s and "," in s:
            s = s.replace(".", "")
            s = s.replace(",", ".")

        # Comma only: could be decimal comma (147,00) or thousands separator
        elif "," in s:
            if re.search(r",\d{1,3}$", s):
                s = s.replace(",", ".")
            else:
                s = s.replace(",", "")

        # Dot only: could be decimal or thousands separator
        elif "." in s:
            parts = s.split(".")
            # If last part is not a typical decimal length, treat dots as thousands separators
            if not (len(parts) == 2 and len(parts[1]) in (1, 2, 3)):
                s = s.replace(".", "")

        # If multiple dots exist, keep the last one as decimal separator
        if s.count(".") > 1:
            p = s.split(".")
            s = "".join(p[:-1]) + "." + p[-1]

        try:
            num = float(s)

            if decimals == 2:
                # CURR fields: ALWAYS 2 decimals (147.00, 2946.85)
                return f"{num:.2f}"

            if decimals == 3:
                # QUAN fields: ALWAYS 3 decimals (5.000)
                return f"{num:.3f}"

            return f"{num:.2f}"
        except Exception:
            return "0.00" if decimals == 2 else "0.000"

    
    # FIXED: Apply correct formatting based on SAP field types
    # Z_TOT_AMT, Z_PRICE, Z_TAX_AMT: CURR (2 decimals)
    sap_grand_total = format_for_sap(grand_total, 2)
    sap_unit_price = format_for_sap(unit_price, 2)
    sap_tax_total = format_for_sap(tax_total, 2)
    
    # Z_QTY: QUAN (3 decimals)
    sap_qty = format_for_sap(qty_str, 3)
    
    # Log the conversion for debugging
    log.info(f"=== SAP FORMATTING ===")
    log.info(f"Original -> SAP Format")
    log.info(f"Z_TOT_AMT: {grand_total} -> {sap_grand_total}")
    log.info(f"Z_PRICE: {unit_price} -> {sap_unit_price}")
    log.info(f"Z_TAX_AMT: {tax_total} -> {sap_tax_total}")
    log.info(f"Z_QTY: {qty_str} -> {sap_qty}")
    log.info(f"Z_CURR: {currency}")
 
    sap_payload = {
        "ZRefNum": get_next_ref_num(),
        "ZInvNum": safe_clean(inv_num),
        "ZInvDate": safe_clean(invoice_date_sap),
        "ZSupName": safe_clean(c_supp),
        "ZTotAmt": safe_clean(sap_grand_total),  # Formatted for SAP CURR field
        "ZCurr": safe_clean(currency),  # Always AED
        "ZQty": safe_clean(sap_qty),    # Formatted for SAP QUAN field
        "ZDesc": safe_clean(clean_desc),
        "ZPrice": safe_clean(sap_unit_price),  # Formatted for SAP CURR field
        "ZTaxAmt": safe_clean(sap_tax_total),  # Formatted for SAP CURR field
    }
    
    log.info(f"Built SAP payload: {json.dumps(sap_payload, indent=2)}")
    return sap_payload
 
# ------------------ SAP posting functions (unchanged) ------------------
def get_sap_session():
    """Create a session with proper SAP authentication"""
    session = requests.Session()
    session.verify = False
    
    # Method 1: Basic Auth
    session.auth = (SAP_USER, SAP_PASS)
    
    # Add SAP headers
    session.headers.update({
        'Accept': 'application/json',
        'Content-Type': 'application/json',
    })
    
    # Add SAP client parameter to session params
    session.params = {'sap-client': SAP_CLIENT}
    
    return session

def post_invoice_to_sap(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Post invoice payload to SAP"""
    base_url = SAP_ODATA_BASE_URL.rstrip('/')
    
    if not base_url.endswith(SAP_ENTITY):
        url = f"{base_url}/{SAP_ENTITY}"
    else:
        url = base_url
    
    log.info(f"🔄 SAP POST URL: {url}")
    log.info(f"🔧 SAP Client: {SAP_CLIENT}")
    log.info(f"👤 SAP User: {SAP_USER}")
    
    try:
        session = get_sap_session()
        
        csrf_token = None
        try:
            csrf_response = session.head(
                url,
                headers={'X-CSRF-Token': 'Fetch'},
                timeout=10
            )
            csrf_token = csrf_response.headers.get('x-csrf-token')
            if csrf_token:
                log.info("✅ CSRF token obtained")
        except Exception as e:
            log.warning(f"CSRF token fetch failed: {e}")
        
        headers = {
            'Content-Type': 'application/json',
            'Accept': 'application/json',
        }
        
        if csrf_token:
            headers['X-CSRF-Token'] = csrf_token
        
        final_url = f"{url}?sap-client={SAP_CLIENT}"
        
        log.info(f"📤 Posting to SAP: {final_url}")
        post_response = session.post(
            final_url,
            headers=headers,
            json=payload,
            timeout=30
        )
        
        result = {
            "status_code": post_response.status_code,
            "success": post_response.status_code in [200, 201, 204],
            "url": final_url,
            "method": "POST"
        }
        
        try:
            result["response"] = post_response.json()
        except:
            result["response"] = post_response.text[:1000] if post_response.text else "Empty response"
        
        log.info(f"📥 SAP Response Status: {post_response.status_code}")
        
        if result["success"]:
            log.info("✅ Invoice successfully posted to SAP")
        else:
            log.error(f"❌ SAP POST failed: {post_response.status_code}")
        
        return result
        
    except Exception as e:
        log.exception("❌ SAP POST failed with exception")
        return {
            "status_code": 500,
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        }

# ------------------ API Endpoints (unchanged) ------------------
@app.get("/health")
def health():
    return {"status": "ok"}
 
@app.post("/debug_ocr_text")
async def debug_ocr_text(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Please upload a PDF file.")
 
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        while True:
            chunk = await file.read(1024 * 1024)
            if not chunk:
                break
            tmp.write(chunk)
        tmp_path = tmp.name
 
    try:
        txt, debug_info = extract_document_text(tmp_path)
        return {"ocr_text_preview": txt[:5000], "debug": debug_info}
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass
 
@app.post("/extract")
async def extract_document(file: UploadFile = File(...)):
    filename = file.filename
    ext = os.path.splitext(filename)[1].lower()
    temp_path = os.path.join(tempfile.gettempdir(), f"{datetime.now().timestamp()}_{filename}")
    
    try:
        with open(temp_path, "wb") as f:
            f.write(await file.read())
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"File write failed: {e}")
 
    text = ""
    debug = {}
    doc_type = "unknown"
    extracted_data = {}
 
    try:
        if ext == ".pdf":
            try:
                gemini_invoices = extract_invoice_with_gemini_pdf(temp_path)

                if gemini_invoices:
                    extracted_data = normalize_gemini_pdf_invoice(gemini_invoices[0])
                    doc_type = "invoice"
                    debug = {"mode": "gemini_pdf"}
                else:
                    raise RuntimeError("Gemini returned empty invoice list")

            except Exception as e:
                log.warning(f"Gemini PDF failed → falling back to OCR: {e}")
                text, debug = extract_document_text(temp_path)
        elif ext in [".docx"]:
            text = extract_text_from_docx(temp_path)
        elif ext in [".doc"]:
            text = extract_text_from_doc(temp_path)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")
 
        doc_type = detect_document_type(text)
        log.info(f"Detected document type: {doc_type}")
        
        if doc_type == "invoice" and extracted_data:
            pass  # Already extracted via Gemini PDF
        elif doc_type in ["invoice", "certificate"]:
            extracted_data = extract_with_gemini(text, doc_type)

        else:
            extracted_data = {"raw_text": text[:5000], "document_type": doc_type}
 
        sap_payload = {}
        sap_response = {}
        if doc_type == "invoice":
            sap_payload = build_sap_invoice_payload(extracted_data)
            sap_response = post_invoice_to_sap(sap_payload)
 
        return JSONResponse({
            "filename": filename,
            "document_type": doc_type,
            "extracted_data": extracted_data,
            "sap_payload": sap_payload,
            "sap_response": sap_response,
            "debug": debug
        })
    finally:
        safe_remove(temp_path)

# ------------------ Approval functions (unchanged) ------------------
def approve_logic(approval_id: str):
    approval = get_approval_status(approval_id)

    if not approval:
        raise HTTPException(status_code=404, detail="Approval not found")

    if approval["status"] != "pending":
        return {
            "sap_ref": approval["sap_payload"].get("ZRefNum", ""),
            "message": "Already processed"
        }

    sap_payload = approval["sap_payload"]

    log.info("Posting invoice to SAP...")
    log.info(json.dumps(sap_payload, indent=2))

    sap_response = post_invoice_to_sap(sap_payload)

    if sap_response.get("status_code") in (200, 201):
        # Get the SAP reference from the payload we sent
        sap_ref = sap_payload.get("ZRefNum", "")
        
        # Pass the SAP reference when updating the status
        update_approval_status(approval_id, "approved", sap_ref)
        
        return {
            "sap_ref": sap_ref,
            "sap_response": sap_response
        }

    update_approval_status(approval_id, "error")
    raise HTTPException(
        status_code=500,
        detail=f"SAP failed: {sap_response}"
    )

def deny_logic(approval_id: str):
    approval = get_approval_status(approval_id)
    if not approval:
        raise HTTPException(status_code=404, detail="Approval not found")

    update_approval_status(approval_id, status="REJECTED")

# ------------------ Additional API endpoints ------------------
@app.post("/extract-handwritten")
async def extract_handwritten_invoice(file: UploadFile = File(...)):
    """Special endpoint for handwritten invoices"""
    filename = file.filename
    ext = os.path.splitext(filename)[1].lower()
    
    if ext not in [".pdf", ".jpg", ".jpeg", ".png"]:
        raise HTTPException(status_code=400, detail="Please upload a PDF or image file.")
    
    temp_path = os.path.join(tempfile.gettempdir(), f"handwritten_{datetime.now().timestamp()}_{filename}")
    
    try:
        with open(temp_path, "wb") as f:
            f.write(await file.read())
        
        # Convert image to PDF if needed
        if ext in [".jpg", ".jpeg", ".png"]:
            from PIL import Image
            img = Image.open(temp_path)
            pdf_path = temp_path + ".pdf"
            img.save(pdf_path, "PDF", resolution=300.0)
            safe_remove(temp_path)
            temp_path = pdf_path
        
        # Extract text with enhanced OCR
        text, debug = enhanced_ocr_pdf_to_text(temp_path)
        
        # Extract data with specialized handwritten extraction
        extracted_data = extract_with_gemini_handwritten(text)
        
        # Build SAP payload
        sap_payload = build_sap_invoice_payload(extracted_data)
        sap_response = post_invoice_to_sap(sap_payload)
        
        return JSONResponse({
            "filename": filename,
            "document_type": "handwritten_invoice",
            "extracted_data": extracted_data,
            "sap_payload": sap_payload,
            "sap_response": sap_response,
            "debug": debug,
            "ocr_text_preview": text[:1000]
        })
        
    finally:
        safe_remove(temp_path)

def process_email_invoice_enhanced(file_path: str, email_subject: str = "", email_from: str = ""):
    """
    Enhanced email invoice processing with multi-page PDF support
    """
    try:
        # Normalize the path for Windows
        file_path = os.path.normpath(file_path)
        filename = os.path.basename(file_path)
        ext = os.path.splitext(filename)[1].lower()
        
        log.info(f"Processing email invoice: {filename}")
        
        approval_ids = []  # Track all approval IDs
        
        if ext == ".pdf":
            try:
                # Open PDF to check page count
                doc = fitz.open(file_path)
                page_count = doc.page_count
                doc.close()
                
                log.info(f"PDF has {page_count} pages")
                
                if page_count > 1:
                    # Multi-page PDF - split and process each page
                    log.info(f"📄 Multi-page PDF detected - splitting into {page_count} individual invoices")
                    
                    # Create temporary directory for split pages
                    temp_dir = tempfile.mkdtemp(prefix="multi_page_")
                    
                    try:
                        # Split PDF into individual pages
                        page_files = split_pdf_by_pages(file_path, temp_dir)
                        
                        # Process each page separately
                        for i, page_file in enumerate(page_files):
                            log.info(f"🔍 Processing invoice {i+1} of {len(page_files)}")
                            
                            # Process single page
                            approval_id = process_single_page_invoice(
                                pdf_path=page_file,
                                page_num=i+1,
                                email_subject=email_subject,
                                email_from=email_from
                            )
                            
                            if approval_id:
                                approval_ids.append(approval_id)
                            
                            # Clean up page file
                            safe_remove(page_file)
                            
                        log.info(f"✅ Processed {len(approval_ids)} invoices from {page_count} pages")
                        
                    finally:
                        # Clean up temp directory
                        try:
                            os.rmdir(temp_dir)
                        except:
                            pass
                            
                else:
                    # Single page PDF - process normally
                    log.info("📄 Single page PDF - processing normally")
                    approval_id = process_single_page_invoice(
                        pdf_path=file_path,
                        page_num=1,
                        email_subject=email_subject,
                        email_from=email_from
                    )
                    
                    if approval_id:
                        approval_ids.append(approval_id)
                        
            except Exception as e:
                log.warning(f"Gemini PDF extraction failed → falling back to OCR: {e}")
                # Fallback to OCR processing for single page
                text, debug = extract_document_text(file_path)
                doc_type = detect_document_type(text)
                
                if doc_type == "invoice":
                    extracted_data = extract_with_gemini(text, doc_type)
                    sap_payload = build_sap_invoice_payload(extracted_data)
                    
                    approval_id = save_for_approval(
                        file_path=file_path,
                        extracted_data=extracted_data,
                        sap_payload=sap_payload,
                        email_subject=email_subject
                    )
                    
                    if approval_id:
                        if send_approval_email(approval_id, extracted_data):
                            log.info(f"✅ Approval email sent for OCR-extracted invoice: {approval_id}")
                            approval_ids.append(approval_id)
                            
        elif ext in [".docx", ".doc"]:
            # Handle Word documents
            if ext == ".docx":
                text = extract_text_from_docx(file_path)
            else:
                text = extract_text_from_doc(file_path)
                
            doc_type = detect_document_type(text)
            
            if doc_type == "invoice":
                extracted_data = extract_with_gemini(text, doc_type)
                sap_payload = build_sap_invoice_payload(extracted_data)
                
                approval_id = save_for_approval(
                    file_path=file_path,
                    extracted_data=extracted_data,
                    sap_payload=sap_payload,
                    email_subject=email_subject
                )
                
                if approval_id:
                    if send_approval_email(approval_id, extracted_data):
                        log.info(f"✅ Approval email sent for Word document: {approval_id}")
                        approval_ids.append(approval_id)
                        
        elif ext in [".jpg", ".jpeg", ".png", ".bmp"]:
            # Convert image to PDF first
            from PIL import Image
            img = Image.open(file_path)
            pdf_path = file_path + ".pdf"
            img.save(pdf_path, "PDF", resolution=300.0)
            
            # Process as single page PDF
            approval_id = process_single_page_invoice(
                pdf_path=pdf_path,
                page_num=1,
                email_subject=email_subject,
                email_from=email_from
            )
            
            if approval_id:
                approval_ids.append(approval_id)
                
            safe_remove(pdf_path)
            
        else:
            log.error(f"Unsupported file type: {ext}")
            return
        
        # Summary log
        if approval_ids:
            log.info(f"🎉 Successfully processed {len(approval_ids)} invoice(s) from {filename}")
            for aid in approval_ids:
                log.info(f"   - Approval ID: {aid}")
        else:
            log.warning(f"No invoices were processed from {filename}")
            
    except Exception as e:
        log.error(f"Error processing email invoice: {e}")
        log.error(traceback.format_exc())

@app.post("/process-email-file")
async def process_email_file(
    file: UploadFile = File(...),
    email_subject: Optional[str] = Form(None),
    email_from: Optional[str] = Form(None)
):
    """
    Endpoint specifically for processing files from email monitoring
    """
    filename = file.filename
    # Use the custom temp directory
    temp_path = os.path.join(CUSTOM_TMP, f"email_{datetime.now().timestamp()}_{filename}")
    
    try:
        # Save uploaded file
        with open(temp_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        log.info(f"Processing email attachment: {filename}")
        log.info(f"Subject: {email_subject}, From: {email_from}")
        
        # Process using enhanced email processing function
        # FIXED: Pass all parameters correctly
        process_email_invoice_enhanced(
            file_path=temp_path, 
            email_subject=email_subject or "", 
            email_from=email_from or ""
        )
        
        return JSONResponse({
            "status": "success",
            "message": "Invoice processed and saved for approval",
            "filename": filename
        })
            
    except Exception as e:
        log.error(f"Error processing email file: {e}")
        return JSONResponse({
            "error": str(e),
            "filename": filename
        }, status_code=500)
            
    finally:
        safe_remove(temp_path)

@app.get("/approve/{approval_id}", response_class=HTMLResponse)
def approve_invoice(approval_id: str):
    try:
        result = approve_logic(approval_id)
        sap_ref = result.get("sap_ref", "")
        
        return HTMLResponse(f"""
        <html>
        <body style="display:flex;justify-content:center;align-items:center;height:100vh;background:#f0fff4">
            <div style="background:white;padding:40px;border-radius:12px;text-align:center;box-shadow:0 10px 30px rgba(0,0,0,0.1);max-width:500px;width:100%">
                <div style="color:#2ecc71;font-size:48px;font-weight:bold;margin-bottom:20px;">✓ APPROVED</div>
                <p style="font-size:18px;color:#555;margin-bottom:30px;">Invoice successfully posted to SAP</p>
                <div style="background:#f8f9fa;padding:20px;border-radius:8px;margin-top:20px;text-align:left;">
                    <p style="margin:10px 0;"><strong> SAP Reference:</strong> <span style="color:#2c3e50;font-weight:bold;">{sap_ref}</span></p>
                    <p style="margin:10px 0;"><strong> Status:</strong> <span style="color:#2ecc71;font-weight:bold;">SUCCESS</span></p>
                    <p style="margin:10px 0;"><strong> Invoice Number:</strong> {result.get("sap_response", {}).get("response", {}).get("d", {}).get("ZInvNum", "N/A")}</p>
                    <p style="margin:10px 0;"><strong> Amount:</strong> {result.get("sap_response", {}).get("response", {}).get("d", {}).get("ZTotAmt", "N/A")} {result.get("sap_response", {}).get("response", {}).get("d", {}).get("ZCurr", "")}</p>
                </div>
                <p style="margin-top:30px;color:#888;font-size:14px;">Confirmation email has been sent to the requester.</p>
            </div>
        </body>
        </html>
        """)
    except HTTPException as e:
        raise e
    except Exception as e:
        error_html = f"""
        <html>
        <body style="display:flex;justify-content:center;align-items:center;height:100vh;background:#fff5f5">
            <div style="background:white;padding:40px;border-radius:12px;text-align:center;box-shadow:0 10px 30px rgba(0,0,0,0.1);max-width:500px;width:100%">
                <div style="color:#e74c3c;font-size:48px;font-weight:bold;margin-bottom:20px;">❌ ERROR</div>
                <p style="font-size:18px;color:#555;margin-bottom:30px;">Failed to approve invoice</p>
                <div style="background:#fff5f5;padding:15px;border-radius:6px;margin-top:20px;">
                    <p style="color:#c0392b;font-size:14px;">Error: {str(e)}</p>
                </div>
            </div>
        </body>
        </html>
        """
        return HTMLResponse(content=error_html, status_code=500)

@app.get("/deny/{approval_id}", response_class=HTMLResponse)
def deny_invoice(approval_id: str):
    try:
        deny_logic(approval_id)
        return HTMLResponse(content="""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Invoice Rejected</title>
            <style>
                body {
                    display: flex;
                    justify-content: center;
                    align-items: center;
                    height: 100vh;
                    background-color: #fff5f5;
                    font-family: Arial, sans-serif;
                }
                .box {
                    text-align: center;
                    padding: 50px;
                    border-radius: 15px;
                    background: white;
                    box-shadow: 0 10px 30px rgba(0,0,0,0.1);
                    max-width: 500px;
                    width: 100%;
                }
                .rejected {
                    color: #e74c3c;
                    font-size: 48px;
                    font-weight: bold;
                    margin-bottom: 20px;
                }
                .sub {
                    margin-top: 15px;
                    font-size: 18px;
                    color: #555;
                }
                .info-box {
                    background: #f8f9fa;
                    padding: 20px;
                    border-radius: 8px;
                    margin-top: 30px;
                    text-align: left;
                }
                .info-box p {
                    margin: 10px 0;
                }
            </style>
        </head>
        <body>
            <div class="box">
                <div class="rejected">✗ REJECTED</div>
                <div class="sub">Invoice has been rejected</div>
                <div class="info-box">
                    <p><strong> Status:</strong> <span style="color:#e74c3c;font-weight:bold;">REJECTED</span></p>
                    <p><strong> Notification:</strong> Confirmation email has been sent to the requester.</p>
                    <p><strong> Action:</strong> No invoice was posted to SAP system.</p>
                </div>
                <p style="margin-top:30px;color:#888;font-size:14px;">The invoice will not be processed for payment.</p>
            </div>
        </body>
        </html>
        """)
    except HTTPException as e:
        raise e
    except Exception as e:
        error_html = f"""
        <html>
        <body style="display:flex;justify-content:center;align-items:center;height:100vh;background:#fff5f5">
            <div style="background:white;padding:40px;border-radius:12px;text-align:center;box-shadow:0 10px 30px rgba(0,0,0,0.1);max-width:500px;width:100%">
                <div style="color:#e74c3c;font-size:48px;font-weight:bold;margin-bottom:20px;">❌ ERROR</div>
                <p style="font-size:18px;color:#555;margin-bottom:30px;">Failed to reject invoice</p>
                <div style="background:#fff5f5;padding:15px;border-radius:6px;margin-top:20px;">
                    <p style="color:#c0392b;font-size:14px;">Error: {str(e)}</p>
                </div>
            </div>
        </body>
        </html>
        """
        return HTMLResponse(content=error_html, status_code=500)

# Also add these approval endpoints if they don't exist:
@app.get("/start-monitoring")
def start_monitoring():
    if start_email_monitoring():
        return {"status": "started", "message": "Email monitoring started"}
    else:
        return {"status": "failed", "message": "Failed to start email monitoring"}

@app.get("/stop-monitoring")
def stop_monitoring():
    stop_email_monitoring()
    return {"status": "stopped", "message": "Email monitoring stopped"}

@app.get("/status/{approval_id}")
def get_status(approval_id: str):
    approval = get_approval_status(approval_id)
    if not approval:
        raise HTTPException(status_code=404, detail="Approval not found")
    return approval

@app.post("/process-bulk-invoices")
async def process_bulk_invoices(
    file: UploadFile = File(...),
    email_subject: Optional[str] = Form(None),
    email_from: Optional[str] = Form(None)
):
    """
    Endpoint for processing PDFs with multiple invoices (one per page)
    """
    filename = file.filename
    temp_path = os.path.join(CUSTOM_TMP, f"bulk_{datetime.now().timestamp()}_{filename}")
    
    try:
        # Save uploaded file
        with open(temp_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        log.info(f"Processing bulk invoices PDF: {filename}")
        
        # Process using enhanced email processing function
        process_email_invoice_enhanced(
            file_path=temp_path, 
            email_subject=email_subject or "", 
            email_from=email_from or ""
        )
        
        return JSONResponse({
            "status": "success",
            "message": "Bulk invoice processing completed",
            "filename": filename,
            "note": "Each page processed as separate invoice with individual approval emails"
        })
            
    except Exception as e:
        log.error(f"Error processing bulk invoices: {e}")
        return JSONResponse({
            "error": str(e),
            "filename": filename
        }, status_code=500)
            
    finally:
        safe_remove(temp_path)

if __name__ == "__main__":
    import uvicorn
    
    init_database()
    os.makedirs("email_attachments", exist_ok=True)
    log.info("=" * 50)
    log.info("Enhanced Invoice Processing System Starting...")
    log.info("✓ Handwritten invoice support enabled")
    log.info(f"✓ Tesseract Arabic support: {ARABIC_AVAILABLE}")
    log.info(f"Server URL: http://{SERVER_HOST}:{SERVER_PORT}")
    log.info("=" * 50)
    log.info("Endpoints:")
    log.info("  POST /extract - Upload any invoice")
    log.info("  POST /extract-handwritten - Upload handwritten invoice (special processing)")
    log.info("  GET /start-monitoring - Start email monitoring")
    log.info("  GET /stop-monitoring - Stop email monitoring")
    log.info("  GET /approve/{id} - Approve invoice")
    log.info("  GET /deny/{id} - Deny invoice")
    log.info("  GET /status/{id} - Check status")
    log.info("=" * 50)
    
    start_email_monitoring()
    
    uvicorn.run(app, host=SERVER_HOST, port=SERVER_PORT)
